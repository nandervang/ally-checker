"""
Word document (ETU format) report generation using python-docx.
"""

from datetime import datetime
from io import BytesIO
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from ..models import AuditData, ReportIssue


def generate_word_report(
    audit_data: AuditData,
    executive_summary: Optional[str],
    locale: str = "sv-SE",
    template: str = "etu-standard"
) -> BytesIO:
    """
    Generate ETU-formatted Word accessibility report.
    
    Args:
        audit_data: Audit results data
        executive_summary: AI-generated summary (optional)
        locale: Language locale
        template: Template type (etu-standard, etu-detailed, etu-summary)
        
    Returns:
        BytesIO buffer containing the Word document
    """
    # Create document
    doc = Document()
    
    # Set document language
    doc.core_properties.language = locale
    
    # Configure default styles
    _setup_styles(doc)
    
    # Generate report sections
    _add_title_page(doc, audit_data, locale)
    doc.add_page_break()
    
    if executive_summary:
        _add_executive_summary(doc, executive_summary, locale)
        doc.add_page_break()
    
    _add_overview_section(doc, audit_data, locale)
    _add_issues_by_principle(doc, audit_data, locale, template)
    _add_compliance_scorecard(doc, audit_data, locale)
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def _setup_styles(doc: Document):
    """Configure document styles for accessibility and ETU branding."""
    styles = doc.styles
    
    # Heading 1 - Title style
    if "Heading 1" not in [s.name for s in styles]:
        h1 = styles.add_style("Heading 1", WD_STYLE_TYPE.PARAGRAPH)
    else:
        h1 = styles["Heading 1"]
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    # Heading 2 - Section headers
    if "Heading 2" not in [s.name for s in styles]:
        h2 = styles.add_style("Heading 2", WD_STYLE_TYPE.PARAGRAPH)
    else:
        h2 = styles["Heading 2"]
    h2.font.size = Pt(18)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 102, 153)  # Medium blue
    
    # Heading 3 - Subsections
    if "Heading 3" not in [s.name for s in styles]:
        h3 = styles.add_style("Heading 3", WD_STYLE_TYPE.PARAGRAPH)
    else:
        h3 = styles["Heading 3"]
    h3.font.size = Pt(14)
    h3.font.bold = True
    
    # Normal text - 18px minimum for accessibility
    normal = styles["Normal"]
    normal.font.size = Pt(12)  # 16px
    normal.font.name = "Calibri"


def _add_title_page(doc: Document, audit_data: AuditData, locale: str):
    """Add ETU-branded title page."""
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(
        "Tillgänglighetsrapport\n" if locale == "sv-SE" else "Accessibility Audit Report\n"
    )
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run(
        "WCAG 2.2 AA Kompatibilitetsbedömning" if locale == "sv-SE" 
        else "WCAG 2.2 AA Compliance Assessment"
    )
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(0, 102, 153)
    
    doc.add_paragraph()  # Spacer
    
    # Metadata table
    metadata_labels = {
        "sv-SE": {
            "url": "Granskad källa:",
            "date": "Granskningsdatum:",
            "type": "Typ av input:",
            "issues": "Totalt antal problem:",
        },
        "en-US": {
            "url": "Audited Source:",
            "date": "Audit Date:",
            "type": "Input Type:",
            "issues": "Total Issues:",
        }
    }
    labels = metadata_labels[locale]
    
    table = doc.add_table(rows=4, cols=2)
    table.style = "Light List Accent 1"
    
    table.rows[0].cells[0].text = labels["url"]
    table.rows[0].cells[1].text = audit_data.url or "HTML Input"
    
    table.rows[1].cells[0].text = labels["date"]
    table.rows[1].cells[1].text = audit_data.created_at.strftime("%Y-%m-%d %H:%M")
    
    table.rows[2].cells[0].text = labels["type"]
    table.rows[2].cells[1].text = audit_data.input_type
    
    table.rows[3].cells[0].text = labels["issues"]
    table.rows[3].cells[1].text = str(audit_data.total_issues)
    
    # Set table font size for accessibility
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)


def _add_executive_summary(doc: Document, summary: str, locale: str):
    """Add AI-generated executive summary section."""
    heading = doc.add_heading(
        "Sammanfattning" if locale == "sv-SE" else "Executive Summary",
        level=1
    )
    
    doc.add_paragraph(summary)


def _add_overview_section(doc: Document, audit_data: AuditData, locale: str):
    """Add overview with WCAG principle breakdown."""
    heading = doc.add_heading(
        "Översikt" if locale == "sv-SE" else "Overview",
        level=1
    )
    
    intro_text = {
        "sv-SE": f"Denna granskning identifierade {audit_data.total_issues} tillgänglighetsproblem "
                 f"organiserade enligt WCAG 2.2 AA:s fyra principer:",
        "en-US": f"This audit identified {audit_data.total_issues} accessibility issues "
                 f"organized by the four WCAG 2.2 AA principles:"
    }
    doc.add_paragraph(intro_text[locale])
    
    # Principle breakdown
    principles = {
        "sv-SE": [
            ("Möjlig att uppfatta", audit_data.perceivable_count),
            ("Hanterbar", audit_data.operable_count),
            ("Begriplig", audit_data.understandable_count),
            ("Robust", audit_data.robust_count),
        ],
        "en-US": [
            ("Perceivable", audit_data.perceivable_count),
            ("Operable", audit_data.operable_count),
            ("Understandable", audit_data.understandable_count),
            ("Robust", audit_data.robust_count),
        ]
    }
    
    for principle_name, count in principles[locale]:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(f"{principle_name}: ").bold = True
        para.add_run(f"{count} " + ("problem" if locale == "sv-SE" else "issues"))


def _add_issues_by_principle(doc: Document, audit_data: AuditData, locale: str, template: str):
    """Add detailed issues grouped by WCAG principle."""
    heading = doc.add_heading(
        "Problem efter princip" if locale == "sv-SE" else "Issues by Principle",
        level=1
    )
    
    # Group issues by principle
    principles = ["Perceivable", "Operable", "Understandable", "Robust"]
    principle_names = {
        "sv-SE": {
            "Perceivable": "Möjlig att uppfatta",
            "Operable": "Hanterbar",
            "Understandable": "Begriplig",
            "Robust": "Robust"
        },
        "en-US": {
            "Perceivable": "Perceivable",
            "Operable": "Operable",
            "Understandable": "Understandable",
            "Robust": "Robust"
        }
    }
    
    for principle in principles:
        principle_issues = [i for i in audit_data.issues if i.wcag_principle == principle]
        if not principle_issues:
            continue
        
        # Add principle heading
        principle_heading = doc.add_heading(
            f"{principle_names[locale][principle]} ({len(principle_issues)})",
            level=2
        )
        
        # Add each issue
        for idx, issue in enumerate(principle_issues, 1):
            _add_issue_detail(doc, issue, idx, locale, template)


def _add_issue_detail(doc: Document, issue: ReportIssue, number: int, locale: str, template: str):
    """Add detailed issue information."""
    # Issue header
    issue_para = doc.add_paragraph()
    issue_para.add_run(f"Problem {number}: " if locale == "sv-SE" else f"Issue {number}: ").bold = True
    issue_para.add_run(issue.success_criterion_name or issue.success_criterion)
    
    # Success criterion
    sc_para = doc.add_paragraph()
    sc_para.add_run("Framgångskriterium: " if locale == "sv-SE" else "Success Criterion: ").bold = True
    sc_para.add_run(issue.success_criterion)
    
    # Severity with color coding
    severity_para = doc.add_paragraph()
    severity_para.add_run("Allvarlighetsgrad: " if locale == "sv-SE" else "Severity: ").bold = True
    severity_run = severity_para.add_run(issue.severity.upper())
    
    # Color code severity
    severity_colors = {
        "critical": RGBColor(153, 0, 0),    # Dark red
        "serious": RGBColor(204, 51, 0),    # Red-orange  
        "moderate": RGBColor(204, 153, 0),  # Orange
        "minor": RGBColor(102, 153, 0),     # Yellow-green
    }
    if issue.severity in severity_colors:
        severity_run.font.color.rgb = severity_colors[issue.severity]
        severity_run.font.bold = True
    
    # Description
    desc_para = doc.add_paragraph()
    desc_para.add_run("Beskrivning: " if locale == "sv-SE" else "Description: ").bold = True
    doc.add_paragraph(issue.description, style="List Bullet")
    
    # Code snippet if available
    if issue.element_snippet and template != "etu-summary":
        code_para = doc.add_paragraph()
        code_para.add_run("Kod: " if locale == "sv-SE" else "Code: ").bold = True
        code_run = code_para.add_run(f"\n{issue.element_snippet}")
        code_run.font.name = "Courier New"
        code_run.font.size = Pt(10)
    
    # Remediation
    remediation_para = doc.add_paragraph()
    remediation_para.add_run("Åtgärd: " if locale == "sv-SE" else "Remediation: ").bold = True
    doc.add_paragraph(issue.remediation, style="List Bullet")
    
    # WCAG reference
    if issue.wcag_reference_url:
        ref_para = doc.add_paragraph()
        ref_para.add_run("WCAG-referens: " if locale == "sv-SE" else "WCAG Reference: ").bold = True
        ref_para.add_run(issue.wcag_reference_url)
    
    doc.add_paragraph()  # Spacer


def _add_compliance_scorecard(doc: Document, audit_data: AuditData, locale: str):
    """Add compliance scorecard summary."""
    doc.add_page_break()
    
    heading = doc.add_heading(
        "Efterlevnadssammanfattning" if locale == "sv-SE" else "Compliance Scorecard",
        level=1
    )
    
    # Determine compliance status
    critical_count = sum(1 for i in audit_data.issues if i.severity == "critical")
    serious_count = sum(1 for i in audit_data.issues if i.severity == "serious")
    
    if critical_count > 0:
        status = "Ej efterlevnad" if locale == "sv-SE" else "Non-Compliant"
        status_color = RGBColor(153, 0, 0)
    elif serious_count > 0:
        status = "Delvis efterlevnad" if locale == "sv-SE" else "Partially Compliant"
        status_color = RGBColor(204, 153, 0)
    elif audit_data.total_issues > 0:
        status = "Huvudsakligen efterlevnad" if locale == "sv-SE" else "Mostly Compliant"
        status_color = RGBColor(102, 153, 0)
    else:
        status = "Fullt efterlevnad" if locale == "sv-SE" else "Fully Compliant"
        status_color = RGBColor(0, 102, 0)
    
    status_para = doc.add_paragraph()
    status_para.add_run("Status: " if locale == "en-US" else "Status: ").bold = True
    status_run = status_para.add_run(status)
    status_run.font.bold = True
    status_run.font.size = Pt(16)
    status_run.font.color.rgb = status_color
    
    doc.add_paragraph()
    
    # Summary table
    table = doc.add_table(rows=5, cols=2)
    table.style = "Light Grid Accent 1"
    
    labels = {
        "sv-SE": ["Totalt problem", "Kritiska", "Allvarliga", "Måttliga", "Mindre"],
        "en-US": ["Total Issues", "Critical", "Serious", "Moderate", "Minor"]
    }
    
    severity_counts = {"critical": 0, "serious": 0, "moderate": 0, "minor": 0}
    for issue in audit_data.issues:
        severity_counts[issue.severity] += 1
    
    table.rows[0].cells[0].text = labels[locale][0]
    table.rows[0].cells[1].text = str(audit_data.total_issues)
    table.rows[1].cells[0].text = labels[locale][1]
    table.rows[1].cells[1].text = str(severity_counts["critical"])
    table.rows[2].cells[0].text = labels[locale][2]
    table.rows[2].cells[1].text = str(severity_counts["serious"])
    table.rows[3].cells[0].text = labels[locale][3]
    table.rows[3].cells[1].text = str(severity_counts["moderate"])
    table.rows[4].cells[0].text = labels[locale][4]
    table.rows[4].cells[1].text = str(severity_counts["minor"])
    
    # Set font sizes
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
    
    # Footer note
    doc.add_paragraph()
    footer_text = {
        "sv-SE": "Detta dokument genererades automatiskt av ETU Tillgänglighetskontroll.",
        "en-US": "This document was automatically generated by ETU Accessibility Checker."
    }
    footer_para = doc.add_paragraph(footer_text[locale])
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
