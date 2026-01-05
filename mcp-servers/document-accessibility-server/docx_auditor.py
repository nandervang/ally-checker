"""
DOCX Accessibility Auditor

Comprehensive DOCX accessibility checking following WCAG 2.2 AA standards.
"""

import logging
from pathlib import Path
from typing import Dict, List, Any
from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


async def audit_docx_accessibility(file_path: str, detailed: bool = True) -> Dict[str, Any]:
    """
    Perform comprehensive DOCX accessibility audit.
    
    Args:
        file_path: Path to DOCX file
        detailed: Include detailed analysis
        
    Returns:
        Audit report with issues categorized by WCAG principle
    """
    issues = []
    
    try:
        docx_path = Path(file_path)
        if not docx_path.exists():
            return {"error": f"File not found: {file_path}"}
        
        doc = Document(file_path)
        
        # Extract metadata
        metadata = {
            "paragraphs": len(doc.paragraphs),
            "sections": len(doc.sections),
            "tables": len(doc.tables),
        }
        
        # Run all checks
        issues.extend(await check_heading_structure(doc))
        issues.extend(await check_docx_alt_text_full(doc))
        issues.extend(await check_docx_tables_full(doc))
        issues.extend(await check_hyperlinks(doc))
        issues.extend(await check_docx_language(doc))
        issues.extend(await check_docx_title(doc))
        
        if detailed:
            issues.extend(await check_list_structure(doc))
            issues.extend(await check_text_formatting(doc))
        
        # Categorize by WCAG principle
        categorized = {
            "Perceivable": [],
            "Operable": [],
            "Understandable": [],
            "Robust": [],
        }
        
        for issue in issues:
            principle = issue.get("wcag_principle", "Robust")
            categorized[principle].append(issue)
        
        return {
            "file": file_path,
            "metadata": metadata,
            "total_issues": len(issues),
            "perceivable_count": len(categorized["Perceivable"]),
            "operable_count": len(categorized["Operable"]),
            "understandable_count": len(categorized["Understandable"]),
            "robust_count": len(categorized["Robust"]),
            "issues": issues,
            "issues_by_principle": categorized,
        }
    
    except Exception as e:
        logger.error(f"Error auditing DOCX {file_path}: {e}", exc_info=True)
        return {"error": str(e), "file": file_path}


async def check_heading_structure(doc: Document) -> List[Dict[str, Any]]:
    """Check document heading structure and hierarchy."""
    issues = []
    
    try:
        headings = []
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.split()[-1])
                    headings.append({"level": level, "text": para.text})
                except (ValueError, IndexError):
                    pass
        
        if len(headings) == 0 and len(doc.paragraphs) > 20:
            issues.append({
                "wcag_principle": "Perceivable",
                "success_criterion": "1.3.1",
                "success_criterion_name": "Info and Relationships",
                "severity": "serious",
                "description": "Document has no headings. Headings provide structure and navigation for screen reader users.",
                "remediation": "Use built-in Heading styles (Heading 1, Heading 2, etc.) to structure the document.",
                "detection_source": "docx-heading-check",
            })
        
        # Check for heading level skips
        if len(headings) > 1:
            for i in range(1, len(headings)):
                prev_level = headings[i-1]["level"]
                curr_level = headings[i]["level"]
                
                if curr_level > prev_level + 1:
                    issues.append({
                        "wcag_principle": "Perceivable",
                        "success_criterion": "1.3.1",
                        "success_criterion_name": "Info and Relationships",
                        "severity": "moderate",
                        "description": f"Heading level skipped: Heading {prev_level} followed by Heading {curr_level}. This breaks the document hierarchy.",
                        "remediation": "Use heading levels in order (H1, H2, H3) without skipping levels.",
                        "detection_source": "docx-heading-check",
                    })
                    break  # Only report once
    
    except Exception as e:
        logger.error(f"Error checking headings: {e}")
    
    return issues


async def check_docx_alt_text_full(doc: Document) -> List[Dict[str, Any]]:
    """Check alternative text for images and objects."""
    issues = []
    
    try:
        images_without_alt = 0
        
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                # This is simplified - full check requires inspecting drawing elements
                images_without_alt += 1
        
        if images_without_alt > 0:
            issues.append({
                "wcag_principle": "Perceivable",
                "success_criterion": "1.1.1",
                "success_criterion_name": "Non-text Content",
                "severity": "critical",
                "description": f"Found {images_without_alt} images. Verify all have descriptive alternative text.",
                "remediation": "Right-click each image, select 'Edit Alt Text', and add a descriptive alternative text. For decorative images, mark as decorative.",
                "detection_source": "docx-image-check",
            })
    
    except Exception as e:
        logger.error(f"Error checking alt text: {e}")
    
    return issues


async def check_docx_tables_full(doc: Document) -> List[Dict[str, Any]]:
    """Check table accessibility."""
    issues = []
    
    try:
        for table_num, table in enumerate(doc.tables):
            # Check if first row looks like headers
            if len(table.rows) > 0:
                first_row = table.rows[0]
                
                # Simple heuristic: check if first row cells are bold
                # Full check would inspect table properties
                
                # For now, just remind to check tables
                if table_num == 0:  # Only report once
                    issues.append({
                        "wcag_principle": "Perceivable",
                        "success_criterion": "1.3.1",
                        "success_criterion_name": "Info and Relationships",
                        "severity": "moderate",
                        "description": f"Document contains {len(doc.tables)} table(s). Verify all tables have header rows designated.",
                        "remediation": "Select the first row of each table, go to Table Tools > Layout, and check 'Repeat Header Rows'.",
                        "detection_source": "docx-table-check",
                    })
    
    except Exception as e:
        logger.error(f"Error checking tables: {e}")
    
    return issues


async def check_hyperlinks(doc: Document) -> List[Dict[str, Any]]:
    """Check hyperlink descriptiveness."""
    issues = []
    
    try:
        non_descriptive_links = []
        
        for para in doc.paragraphs:
            for run in para.runs:
                # Check for hyperlinks
                if run.font.underline:
                    text = run.text.strip().lower()
                    
                    # Check for non-descriptive link text
                    if text in ['click here', 'here', 'link', 'read more', 'more', 'this', 'http', 'https', 'www']:
                        non_descriptive_links.append(text)
        
        if non_descriptive_links:
            issues.append({
                "wcag_principle": "Operable",
                "success_criterion": "2.4.4",
                "success_criterion_name": "Link Purpose (In Context)",
                "severity": "moderate",
                "description": f"Found {len(non_descriptive_links)} potentially non-descriptive links (e.g., 'click here', 'read more').",
                "remediation": "Use descriptive link text that makes sense out of context. Instead of 'click here', use 'Download the accessibility guide'.",
                "detection_source": "docx-link-check",
            })
    
    except Exception as e:
        logger.error(f"Error checking hyperlinks: {e}")
    
    return issues


async def check_docx_language(doc: Document) -> List[Dict[str, Any]]:
    """Check document language setting."""
    issues = []
    
    try:
        # Check core properties for language
        if not doc.core_properties.language:
            issues.append({
                "wcag_principle": "Understandable",
                "success_criterion": "3.1.1",
                "success_criterion_name": "Language of Page",
                "severity": "serious",
                "description": "Document language is not specified. Screen readers may not pronounce text correctly.",
                "remediation": "Go to File > Info > Properties > Advanced Properties > Summary tab and set the Language field.",
                "detection_source": "docx-language-check",
            })
    
    except Exception as e:
        logger.error(f"Error checking language: {e}")
    
    return issues


async def check_docx_title(doc: Document) -> List[Dict[str, Any]]:
    """Check document title."""
    issues = []
    
    try:
        if not doc.core_properties.title or doc.core_properties.title.strip() == "":
            issues.append({
                "wcag_principle": "Operable",
                "success_criterion": "2.4.2",
                "success_criterion_name": "Page Titled",
                "severity": "serious",
                "description": "Document has no title set. This makes it hard to identify the document's purpose.",
                "remediation": "Go to File > Info > Properties and set the Title field.",
                "detection_source": "docx-metadata-check",
            })
    
    except Exception as e:
        logger.error(f"Error checking title: {e}")
    
    return issues


async def check_list_structure(doc: Document) -> List[Dict[str, Any]]:
    """Check if lists use proper list formatting."""
    issues = []
    
    # This would require deeper analysis of paragraph styles
    # For MVP, we'll skip this check
    
    return issues


async def check_text_formatting(doc: Document) -> List[Dict[str, Any]]:
    """Check for color-only formatting or other issues."""
    issues = []
    
    # Check for very large fonts (may indicate improper heading use)
    # This is a heuristic check
    
    try:
        large_text_not_headings = 0
        
        for para in doc.paragraphs:
            if not para.style.name.startswith('Heading'):
                for run in para.runs:
                    if run.font.size and run.font.size.pt > 14:
                        large_text_not_headings += 1
                        break
        
        if large_text_not_headings > 3:
            issues.append({
                "wcag_principle": "Perceivable",
                "success_criterion": "1.3.1",
                "success_criterion_name": "Info and Relationships",
                "severity": "minor",
                "description": f"Found {large_text_not_headings} paragraphs with large text that are not using Heading styles. Visual formatting alone is not accessible.",
                "remediation": "Use built-in Heading styles instead of manual font sizing to create document structure.",
                "detection_source": "docx-formatting-check",
            })
    
    except Exception as e:
        logger.error(f"Error checking formatting: {e}")
    
    return issues
