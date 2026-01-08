"""
DOCX-specific accessibility tools.
"""

from typing import Dict, Any, List
from docx import Document


def extract_docx_structure(file_path: str) -> Dict[str, Any]:
    """Extract DOCX structure information."""
    try:
        doc = Document(file_path)
        
        headings = []
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.split()[-1])
                    headings.append({
                        "level": level,
                        "text": para.text[:100],  # First 100 chars
                    })
                except (ValueError, IndexError):
                    pass
        
        return {
            "paragraphs": len(doc.paragraphs),
            "sections": len(doc.sections),
            "tables": len(doc.tables),
            "headings": headings,
            "heading_count": len(headings),
        }
    
    except Exception as e:
        return {"error": str(e)}


def check_docx_headings(file_path: str) -> Dict[str, Any]:
    """Check heading structure."""
    try:
        doc = Document(file_path)
        headings = []
        
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.split()[-1])
                    headings.append({"level": level, "text": para.text})
                except (ValueError, IndexError):
                    pass
        
        issues = []
        if len(headings) == 0:
            issues.append("No headings found")
        
        # Check for level skips
        for i in range(1, len(headings)):
            if headings[i]["level"] > headings[i-1]["level"] + 1:
                issues.append(f"Heading level skip detected: H{headings[i-1]['level']} to H{headings[i]['level']}")
        
        return {
            "heading_count": len(headings),
            "headings": headings,
            "issues": issues,
        }
    
    except Exception as e:
        return {"error": str(e)}


def check_docx_alt_text(file_path: str) -> Dict[str, Any]:
    """Check alt text for images."""
    try:
        doc = Document(file_path)
        
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
        
        return {
            "image_count": image_count,
            "message": f"Found {image_count} images - verify all have descriptive alt text",
        }
    
    except Exception as e:
        return {"error": str(e)}


def check_docx_tables(file_path: str) -> Dict[str, Any]:
    """Check table structure."""
    try:
        doc = Document(file_path)
        
        tables_info = []
        for table in doc.tables:
            tables_info.append({
                "rows": len(table.rows),
                "columns": len(table.columns) if table.rows else 0,
            })
        
        return {
            "table_count": len(tables_info),
            "tables": tables_info,
            "message": "Verify all tables have header rows designated",
        }
    
    except Exception as e:
        return {"error": str(e)}


def check_docx_hyperlinks(file_path: str) -> Dict[str, Any]:
    """Check hyperlink descriptiveness."""
    try:
        doc = Document(file_path)
        
        potential_issues = []
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.underline:
                    text = run.text.strip().lower()
                    if text in ['click here', 'here', 'link', 'read more']:
                        potential_issues.append(text)
        
        return {
            "non_descriptive_links": len(potential_issues),
            "examples": list(set(potential_issues))[:5],
            "message": "Use descriptive link text instead of 'click here' or 'read more'",
        }
    
    except Exception as e:
        return {"error": str(e)}
