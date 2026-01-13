
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_test_docx():
    doc = Document()
    doc.core_properties.title = ""  # Missing title
    # Missing language (cannot easily unset if default, but we can check if we set it)
    
    # Add heading skip
    # Level 1 = Heading 1
    doc.add_heading('Heading 1', 1) 
    doc.add_paragraph('Some text.')
    
    # Level 3 = Heading 3 (skips 2)
    doc.add_heading('Jumped to Heading 3', 3) 
    
    # Add non-descriptive link
    p = doc.add_paragraph('Click ')
    r = p.add_run('here')
    r.font.underline = True
    r.font.color.rgb = RGBColor(0, 0, 255)
    p.add_run(' for more info.')
    
    # Add valid link
    p2 = doc.add_paragraph('Download the ')
    r2 = p2.add_run('accessibility report')
    r2.font.underline = True
    r2.font.color.rgb = RGBColor(0, 0, 255)
    p2.add_run('.')

    # Add table with no header row styling
    table = doc.add_table(rows=3, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Value'
    
    # Add large text not heading
    p_large = doc.add_paragraph('This looks like a heading but is just big text')
    run_large = p_large.runs[0]
    run_large.font.size = Pt(24)

    doc.save('test_bad.docx')
    print("Created test_bad.docx")

if __name__ == "__main__":
    create_test_docx()
